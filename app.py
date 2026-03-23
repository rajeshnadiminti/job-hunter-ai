"""
Job Hunter Pro — Multi-AI Career Command Center
A local Python/Flask application for AI-powered job hunting.
Run: python app.py
"""

import os, json, csv, io, time, base64, re, uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file, Response
import requests as http_req
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20MB upload limit

DATA_FILE = Path(__file__).parent / "data.json"
UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# ═══════════════════════════════════════════════════════════
# DATA PERSISTENCE — simple JSON file
# ═══════════════════════════════════════════════════════════
def load_data():
    if DATA_FILE.exists():
        return json.loads(DATA_FILE.read_text())
    return {"jobs": [], "settings": {"provider": "anthropic", "keys": {}, "models": {}}, "resume": ""}

def save_data(data):
    DATA_FILE.write_text(json.dumps(data, indent=2))

# ═══════════════════════════════════════════════════════════
# AI PROVIDERS
# ═══════════════════════════════════════════════════════════
PROVIDERS = {
    "anthropic": {
        "name": "Anthropic", "short": "Claude",
        "models": [
            {"id": "claude-opus-4-6-20260205", "label": "Claude Opus 4.6", "tag": "Most Intelligent"},
            {"id": "claude-sonnet-4-6-20260217", "label": "Claude Sonnet 4.6", "tag": "Best Value"},
            {"id": "claude-opus-4-5-20251120", "label": "Claude Opus 4.5", "tag": "Previous Opus"},
            {"id": "claude-sonnet-4-5-20241022", "label": "Claude Sonnet 4.5", "tag": "Previous Sonnet"},
            {"id": "claude-sonnet-4-20250514", "label": "Claude Sonnet 4", "tag": "Legacy"},
            {"id": "claude-haiku-4-5-20251001", "label": "Claude Haiku 4.5", "tag": "Fastest"},
        ],
        "default": "claude-sonnet-4-6-20260217",
    },
    "openai": {
        "name": "OpenAI", "short": "GPT",
        "models": [
            {"id": "gpt-5.4", "label": "GPT-5.4", "tag": "Flagship"},
            {"id": "gpt-5.4-mini", "label": "GPT-5.4 Mini", "tag": "Recommended"},
            {"id": "gpt-5.4-nano", "label": "GPT-5.4 Nano", "tag": "Fastest"},
            {"id": "gpt-5.2", "label": "GPT-5.2", "tag": "Previous Gen"},
            {"id": "o3", "label": "o3", "tag": "Reasoning"},
            {"id": "o4-mini", "label": "o4-mini", "tag": "Fast Reasoning"},
            {"id": "gpt-4o", "label": "GPT-4o", "tag": "Legacy"},
            {"id": "gpt-4.1", "label": "GPT-4.1", "tag": "1M Context"},
        ],
        "default": "gpt-5.4-mini",
    },
    "gemini": {
        "name": "Google", "short": "Gemini",
        "models": [
            {"id": "gemini-3.1-pro-preview", "label": "Gemini 3.1 Pro", "tag": "Latest"},
            {"id": "gemini-3-flash-preview", "label": "Gemini 3 Flash", "tag": "Fast + Smart"},
            {"id": "gemini-2.5-pro", "label": "Gemini 2.5 Pro", "tag": "Stable"},
            {"id": "gemini-2.5-flash", "label": "Gemini 2.5 Flash", "tag": "Best Value"},
            {"id": "gemini-2.5-flash-lite", "label": "Gemini 2.5 Flash Lite", "tag": "Cheapest"},
        ],
        "default": "gemini-2.5-flash",
    },
}

def call_ai(provider, api_key, model, system_prompt, user_message, web_search=False):
    """Unified AI call — works with all three providers."""
    if provider == "anthropic":
        if not api_key:
            raise ValueError("Anthropic API key required")
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        }
        body = {
            "model": model,
            "max_tokens": 4096,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_message}],
        }
        if web_search:
            body["tools"] = [{"type": "web_search_20250305", "name": "web_search"}]
        r = http_req.post("https://api.anthropic.com/v1/messages", headers=headers, json=body, timeout=120)
        r.raise_for_status()
        d = r.json()
        return "\n".join(b["text"] for b in d.get("content", []) if b.get("text"))

    elif provider == "openai":
        if not api_key:
            raise ValueError("OpenAI API key required")
        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
        body = {
            "model": model,
            "max_tokens": 4096,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
        }
        r = http_req.post("https://api.openai.com/v1/chat/completions", headers=headers, json=body, timeout=120)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"]

    elif provider == "gemini":
        if not api_key:
            raise ValueError("Gemini API key required")
        body = {
            "contents": [{"parts": [{"text": user_message}]}],
            "systemInstruction": {"parts": [{"text": system_prompt}]},
            "generationConfig": {"maxOutputTokens": 4096},
        }
        if web_search:
            body["tools"] = [{"googleSearch": {}}]
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        r = http_req.post(url, json=body, timeout=120)
        r.raise_for_status()
        d = r.json()
        parts = d.get("candidates", [{}])[0].get("content", {}).get("parts", [])
        return "\n".join(p.get("text", "") for p in parts)

    raise ValueError(f"Unknown provider: {provider}")

# ═══════════════════════════════════════════════════════════
# FILE PROCESSING
# ═══════════════════════════════════════════════════════════
def extract_pdf_text(filepath):
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts)

def extract_docx_text(filepath):
    doc = DocxDocument(filepath)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def generate_docx(text, title="Document"):
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # Headings
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:])
            p.style = doc.styles["Heading 2"]
        elif stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Heading 1"]
        elif re.match(r"^[A-Z][A-Z\s&|/\-]{4,}$", stripped) or stripped.endswith(":"):
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("• "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            # Handle **bold** in text
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════
# ROUTES — Pages
# ═══════════════════════════════════════════════════════════
@app.route("/")
def index():
    data = load_data()
    return render_template("index.html", providers=PROVIDERS, data=data)

# ═══════════════════════════════════════════════════════════
# ROUTES — Settings
# ═══════════════════════════════════════════════════════════
@app.post("/api/settings")
def save_settings():
    data = load_data()
    body = request.json
    data["settings"] = {
        "provider": body.get("provider", "anthropic"),
        "keys": body.get("keys", {}),
        "models": body.get("models", {}),
    }
    save_data(data)
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
# ROUTES — Resume
# ═══════════════════════════════════════════════════════════
@app.post("/api/resume/upload")
def upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    ext = f.filename.rsplit(".", 1)[-1].lower()
    filepath = UPLOAD_DIR / f"resume.{ext}"
    f.save(filepath)
    try:
        if ext == "pdf":
            text = extract_pdf_text(filepath)
        elif ext in ("docx", "doc"):
            text = extract_docx_text(filepath)
        else:
            text = filepath.read_text(errors="replace")
        data = load_data()
        data["resume"] = text
        save_data(data)
        return jsonify({"text": text, "filename": f.filename, "size": os.path.getsize(filepath)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.post("/api/resume/save")
def save_resume_text():
    data = load_data()
    data["resume"] = request.json.get("text", "")
    save_data(data)
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
# ROUTES — AI Actions
# ═══════════════════════════════════════════════════════════
def get_ai_config():
    data = load_data()
    s = data["settings"]
    provider = s.get("provider", "anthropic")
    key = s.get("keys", {}).get(provider, "")
    model = s.get("models", {}).get(provider, PROVIDERS[provider]["default"])
    return provider, key, model, data.get("resume", "")

@app.post("/api/ai/search")
def ai_search():
    body = request.json
    provider, key, model, resume = get_ai_config()
    platforms = body.get("platforms", "LinkedIn, Indeed")
    query = body.get("query", "")
    location = body.get("location", "")
    sys_prompt = f"Job search assistant. Find real current job listings. For each job provide: Title, Company, Location, Salary, Direct Apply Link (full URL), Platform, Contact/Recruiter email if visible. Search across: {platforms}."
    user_msg = f"Find jobs matching: {query}" + (f" in {location}" if location else "") + f". Search on {platforms}. Return real current postings with direct apply links and salary info."
    if resume:
        user_msg += f"\n\nCandidate profile (match jobs to this):\n{resume[:2000]}"
    try:
        result = call_ai(provider, key, model, sys_prompt, user_msg, web_search=True)
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.post("/api/ai/parse-jobs")
def ai_parse_jobs():
    body = request.json
    provider, key, model, _ = get_ai_config()
    text = body.get("text", "")[:4000]
    sys_prompt = 'Extract job listings into a JSON array. Return ONLY valid JSON. Each object: {"title":"","company":"","location":"","salary":"","link":"","platform":"","contact":""}. Use "" for missing fields.'
    try:
        result = call_ai(provider, key, model, sys_prompt, f"Extract jobs:\n{text}", web_search=False)
        cleaned = re.sub(r"```json\s*|```\s*", "", result).strip()
        match = re.search(r"\[[\s\S]*\]", cleaned)
        if match:
            jobs = json.loads(match.group())
            return jsonify({"jobs": jobs})
        return jsonify({"jobs": []})
    except Exception as e:
        return jsonify({"error": str(e), "jobs": []}), 200

@app.post("/api/ai/tool")
def ai_tool():
    body = request.json
    provider, key, model, resume = get_ai_config()
    action = body.get("action", "")
    job = body.get("job", {})
    title = job.get("title", "")
    company = job.get("company", "")
    location = job.get("location", "")
    notes = job.get("notes", "")[:600]

    prompts = {
        "tailor": (
            "Expert resume writer. Tailor this resume for the target job. Use ATS keywords, quantify achievements. Use clear sections: Summary, Experience, Skills, Education. Use ## for headings and - for bullet points.",
            f"TARGET: {title} at {company}\n{('JD: ' + notes) if notes else ''}\n\nRESUME:\n{resume[:6000] or '(No resume loaded — give general advice)'}",
            False
        ),
        "cover": (
            "Expert cover letter writer. Write a personalized cover letter with business format: date, salutation, 3-4 paragraphs, sign-off. Be specific about the company.",
            f"JOB: {title} at {company}, {location}\n{('JD: ' + notes) if notes else ''}\n\nRESUME:\n{resume[:4000] or '(General)'}",
            False
        ),
        "summary": (
            "Career branding expert. Create: 1) Professional summary 2) Elevator pitch 3) LinkedIn headline 4) Short bio. All tailored to the target role.",
            f"TARGET: {title} at {company}\nRESUME:\n{resume[:4000] or '(General)'}",
            False
        ),
        "interview": (
            "Interview coach. Generate 10 likely questions (behavioral, technical, situational) with answer frameworks. Include 5 questions the candidate should ask.",
            f"ROLE: {title} at {company}\n{('JD: ' + notes) if notes else ''}\nBACKGROUND:\n{resume[:3000] or '(General)'}",
            False
        ),
        "salary": (
            "Salary negotiation expert. Provide salary benchmarks, ranges, and negotiation strategies with scripts.",
            f"Research salary: {title} at {company} in {location or 'US'}. Give benchmarks and negotiation tips.",
            True
        ),
        "network": (
            "Networking expert. Create: 1) LinkedIn connection request 2) Cold email template 3) Follow-up sequence.",
            f"ROLE: {title} at {company}\nContact: {job.get('contact', 'unknown')}",
            False
        ),
    }

    if action not in prompts:
        return jsonify({"error": "Unknown action"}), 400

    sys_p, usr_p, ws = prompts[action]
    try:
        result = call_ai(provider, key, model, sys_p, usr_p, web_search=ws)
        return jsonify({"result": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ═══════════════════════════════════════════════════════════
# ROUTES — Jobs CRUD
# ═══════════════════════════════════════════════════════════
@app.get("/api/jobs")
def get_jobs():
    return jsonify(load_data().get("jobs", []))

@app.post("/api/jobs")
def add_job():
    data = load_data()
    job = request.json
    job["id"] = str(uuid.uuid4())[:8]
    job["status"] = job.get("status", "saved")
    job["dateAdded"] = datetime.now().strftime("%Y-%m-%d")
    job["starred"] = False
    data["jobs"].append(job)
    save_data(data)
    return jsonify(job)

@app.put("/api/jobs/<job_id>")
def update_job(job_id):
    data = load_data()
    updates = request.json
    for j in data["jobs"]:
        if j["id"] == job_id:
            j.update(updates)
            break
    save_data(data)
    return jsonify({"ok": True})

@app.delete("/api/jobs/<job_id>")
def delete_job(job_id):
    data = load_data()
    data["jobs"] = [j for j in data["jobs"] if j["id"] != job_id]
    save_data(data)
    return jsonify({"ok": True})

# ═══════════════════════════════════════════════════════════
# ROUTES — Export
# ═══════════════════════════════════════════════════════════
@app.get("/api/export/csv")
def export_csv():
    data = load_data()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Title", "Company", "Location", "Salary", "Platform", "Status", "Link", "Contact", "Date", "Notes"])
    for j in data["jobs"]:
        writer.writerow([j.get(k, "") for k in ["title", "company", "location", "salary", "platform", "status", "link", "contact", "dateAdded", "notes"]])
    output.seek(0)
    return Response(output.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": f"attachment; filename=jobs_{datetime.now().strftime('%Y-%m-%d')}.csv"})

@app.post("/api/export/docx")
def export_docx():
    body = request.json
    text = body.get("text", "")
    filename = body.get("filename", "document")
    buf = generate_docx(text, filename)
    return send_file(buf, as_attachment=True, download_name=f"{filename}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ═══════════════════════════════════════════════════════════
# RUN
# ═══════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("  JOB HUNTER PRO — Multi-AI Career Command Center")
    print("=" * 60)
    print(f"\n  Open in browser:  http://localhost:5000")
    print(f"  Data stored in:   {DATA_FILE}")
    print(f"  Uploads dir:      {UPLOAD_DIR}")
    print(f"\n  Press Ctrl+C to stop\n")
    app.run(debug=True, port=5000)
